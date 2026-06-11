#!/usr/bin/env python3
"""
Convert Markdown to DOCX with customizable formatting.

Usage:
    python3 tools/md_to_docx.py input.md [output.docx] [--style style.json]

Style config: JSON file with format spec. See assets/default-style.json.
"""

import argparse
import io
import json
import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ASSETS_DIR = Path(__file__).parent.parent / "assets"


def _mermaid_js_src() -> str:
    local = Path(__file__).parent.parent / "node_modules" / "mermaid" / "dist" / "mermaid.min.js"
    if local.exists():
        return local.as_uri()
    return "https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"


def render_mermaid_png(mermaid_code: str) -> bytes | None:
    """Render Mermaid diagram to PNG via Playwright. Returns None if unavailable."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return None

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<script src="{_mermaid_js_src()}"></script>
<style>body{{margin:0;background:white}}.mermaid{{display:inline-block}}</style>
</head><body>
<pre class="mermaid">{mermaid_code}</pre>
<script>mermaid.initialize({{startOnLoad:true}});</script>
</body></html>"""

    tmp = Path(tempfile.mktemp(suffix=".html"))
    try:
        tmp.write_text(html, encoding="utf-8")
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(tmp.as_uri())
            page.wait_for_function(
                "document.querySelector('pre.mermaid svg') !== null",
                timeout=10000,
            )
            el = page.query_selector("pre.mermaid")
            png = el.screenshot()
            browser.close()
        return png
    except Exception:
        return None
    finally:
        tmp.unlink(missing_ok=True)

# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Apply selective cell borders. Each side: dict {val, sz, color} or None to skip."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side_name, cfg in [("top", top), ("bottom", bottom),
                            ("left", left), ("right", right)]:
        if cfg is None:
            continue
        el = OxmlElement(f"w:{side_name}")
        el.set(qn("w:val"), cfg.get("val", "single"))
        if cfg.get("val") != "nil":
            el.set(qn("w:sz"), str(cfg.get("sz", 4)))
            el.set(qn("w:color"), cfg.get("color", "000000").lstrip("#"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def apply_run_format(run, font_name: str, font_en: str, size_pt: float,
                     bold: bool = False, italic: bool = False,
                     color: str | None = None):
    run.font.name = font_en
    run._r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = hex_to_rgb(color)


def set_paragraph_format(para, space_before: float = 0, space_after: float = 0,
                          line_spacing: float | None = None,
                          first_line_indent: float | None = None,
                          left_indent: float | None = None,
                          align: str = "left"):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)


# ── Inline markdown parser ────────────────────────────────────────────────────
# Handles: **bold**, *italic*, `code`, ~~strikethrough~~, [text](url)

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*"
    r"|__(?P<bold2>.+?)__"
    r"|\*(?P<italic>.+?)\*"
    r"|_(?P<italic2>.+?)_"
    r"|`(?P<code>.+?)`"
    r"|~~(?P<strike>.+?)~~"
    r"|\[(?P<link_text>[^\]]+)\]\([^\)]+\))"
)


def add_inline(paragraph, text: str, base_font: str, base_font_en: str,
               base_size: float, base_bold: bool = False,
               base_italic: bool = False, base_color: str | None = None):
    """Parse inline markdown and add runs to paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # plain text before match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            apply_run_format(run, base_font, base_font_en, base_size,
                             base_bold, base_italic, base_color)

        if m.group("bold") or m.group("bold2"):
            content = m.group("bold") or m.group("bold2")
            run = paragraph.add_run(content)
            apply_run_format(run, base_font, base_font_en, base_size,
                             bold=True, italic=base_italic, color=base_color)
        elif m.group("italic") or m.group("italic2"):
            content = m.group("italic") or m.group("italic2")
            run = paragraph.add_run(content)
            apply_run_format(run, base_font, base_font_en, base_size,
                             bold=base_bold, italic=True, color=base_color)
        elif m.group("code"):
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Courier New"
            run._r.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(base_size - 1)
        elif m.group("strike"):
            run = paragraph.add_run(m.group("strike"))
            apply_run_format(run, base_font, base_font_en, base_size,
                             base_bold, base_italic, base_color)
            run.font.strike = True
        elif m.group("link_text"):
            run = paragraph.add_run(m.group("link_text"))
            apply_run_format(run, base_font, base_font_en, base_size,
                             base_bold, base_italic, "1155CC")
            run.font.underline = True

        pos = m.end()

    # trailing plain text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        apply_run_format(run, base_font, base_font_en, base_size,
                         base_bold, base_italic, base_color)


# ── Block-level parser ────────────────────────────────────────────────────────

def parse_md_blocks(md_text: str) -> list[dict]:
    """
    Tokenise markdown into a flat list of block dicts:
      {type: heading|paragraph|code|blockquote|hr|list_item|blank}
    """
    lines = md_text.splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # HTML comment — skip
        if re.match(r"^\s*<!--.*-->\s*$", line) or line.strip().startswith("<!--"):
            # skip until end of multi-line comment
            if "-->" not in line:
                while i < len(lines) and "-->" not in lines[i]:
                    i += 1
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            text = "\n".join(code_lines)
            if lang == "mermaid":
                blocks.append({"type": "mermaid", "text": text})
            else:
                blocks.append({"type": "code", "lang": lang, "text": text})
            i += 1
            continue

        # ATX headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level,
                           "text": m.group(2).strip()})
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}\s*$", line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Blockquote (collect consecutive > lines)
        if line.startswith(">"):
            bq_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                bq_lines.append(lines[i].lstrip(">").strip())
                i += 1
            blocks.append({"type": "blockquote",
                           "text": " ".join(bq_lines)})
            continue

        # Unordered list item
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            blocks.append({"type": "list_item", "ordered": False,
                           "indent": indent, "text": m.group(2).strip()})
            i += 1
            continue

        # Ordered list item
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            blocks.append({"type": "list_item", "ordered": True,
                           "indent": indent, "text": m.group(2).strip()})
            i += 1
            continue

        # Markdown pipe table
        if re.match(r"^\s*\|", line):
            table_lines = []
            while i < len(lines) and re.match(r"^\s*\|", lines[i]):
                table_lines.append(lines[i])
                i += 1
            # Parse header, separator, rows
            def split_row(r):
                return [c.strip() for c in r.strip().strip("|").split("|")]
            header = split_row(table_lines[0])
            rows = []
            for tl in table_lines[2:]:  # skip separator line
                rows.append(split_row(tl))
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        # Standalone image: ![alt](path)
        m = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if m:
            blocks.append({"type": "image", "alt": m.group(1), "src": m.group(2)})
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            blocks.append({"type": "blank"})
            i += 1
            continue

        # Paragraph — collect until blank or block-level element
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if (next_line.strip() == ""
                    or next_line.startswith("#")
                    or next_line.startswith(">")
                    or next_line.startswith("```")
                    or re.match(r"^[-*_]{3,}\s*$", next_line)
                    or re.match(r"^\s*\|", next_line)
                    or re.match(r"^(\s*)[-*+]\s+", next_line)
                    or re.match(r"^(\s*)\d+\.\s+", next_line)):
                break
            para_lines.append(next_line)
            i += 1
        blocks.append({"type": "paragraph",
                       "text": " ".join(para_lines)})

    return blocks


# ── Document builder ──────────────────────────────────────────────────────────

def build_docx(blocks: list[dict], style: dict, base_dir: Path | None = None) -> Document:
    doc = Document()

    # Page margins
    pg = style["page"]
    for section in doc.sections:
        section.top_margin = Cm(pg["top_cm"])
        section.bottom_margin = Cm(pg["bottom_cm"])
        section.left_margin = Cm(pg["left_cm"])
        section.right_margin = Cm(pg["right_cm"])

    body = style["body"]
    headings = style["headings"]
    code_cfg = style["code_block"]
    bq_cfg = style["blockquote"]
    tbl_cfg = style.get("table", {})

    # One char ≈ font_size pt; first-line indent in pt
    char_width_pt = body["size_pt"]
    first_indent_pt = char_width_pt * body.get("first_line_indent_chars", 2)

    list_counters: dict[int, int] = {}

    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue

        elif btype == "hr":
            para = doc.add_paragraph()
            run = para.add_run("─" * 40)
            run.font.color.rgb = hex_to_rgb("AAAAAA")
            set_paragraph_format(para, space_before=4, space_after=4,
                                 align="center")

        elif btype == "heading":
            level = min(block["level"], 4)
            key = f"h{level}"
            cfg = headings.get(key, headings["h4"])
            para = doc.add_paragraph()
            add_inline(para, block["text"],
                       cfg["font"], cfg["font_en"], cfg["size_pt"],
                       base_bold=cfg.get("bold", False),
                       base_color=cfg.get("color"))
            set_paragraph_format(para,
                                 space_before=cfg.get("space_before_pt", 12),
                                 space_after=cfg.get("space_after_pt", 6),
                                 align=cfg.get("align", "left"))

        elif btype == "paragraph":
            para = doc.add_paragraph()
            add_inline(para, block["text"],
                       body["font"], body["font_en"], body["size_pt"])
            set_paragraph_format(para,
                                 space_before=body.get("space_before_pt", 0),
                                 space_after=body.get("space_after_pt", 6),
                                 line_spacing=body.get("line_spacing_pt"),
                                 first_line_indent=first_indent_pt)

        elif btype == "code":
            para = doc.add_paragraph()
            run = para.add_run(block["text"])
            run.font.name = code_cfg["font"]
            run._r.rPr.rFonts.set(qn("w:eastAsia"), code_cfg["font"])
            run.font.size = Pt(code_cfg["size_pt"])
            set_paragraph_format(para, space_before=6, space_after=6,
                                 left_indent=1.0)

        elif btype == "image":
            src = block["src"]
            img_path = Path(src) if Path(src).is_absolute() else (base_dir / src if base_dir else Path(src))
            para = doc.add_paragraph()
            set_paragraph_format(para, space_before=6, space_after=6, align="center")
            run = para.add_run()
            if img_path.exists():
                run.add_picture(str(img_path), width=Cm(14))
            else:
                run.italic = True
                run.text = f"[图片: {block['alt'] or src}]"
                apply_run_format(run, body["font"], body["font_en"], body["size_pt"],
                                 italic=True, color="888888")

        elif btype == "mermaid":
            png = render_mermaid_png(block["text"])
            if png:
                para = doc.add_paragraph()
                set_paragraph_format(para, space_before=6, space_after=6,
                                     align="center")
                run = para.add_run()
                run.add_picture(io.BytesIO(png), width=Cm(14))
            else:
                # Playwright unavailable — render as fenced code block
                para = doc.add_paragraph()
                run = para.add_run("```mermaid\n" + block["text"] + "\n```")
                run.font.name = code_cfg["font"]
                run._r.rPr.rFonts.set(qn("w:eastAsia"), code_cfg["font"])
                run.font.size = Pt(code_cfg["size_pt"])
                set_paragraph_format(para, space_before=6, space_after=6,
                                     left_indent=1.0)

        elif btype == "blockquote":
            para = doc.add_paragraph()
            add_inline(para, block["text"],
                       bq_cfg["font"], bq_cfg["font_en"], bq_cfg["size_pt"],
                       base_color=bq_cfg.get("color"))
            set_paragraph_format(para, space_before=4, space_after=4,
                                 left_indent=bq_cfg.get("left_indent_cm", 1.0))

        elif btype == "list_item":
            indent = block.get("indent", 0)
            ordered = block.get("ordered", False)

            if ordered:
                list_counters.setdefault(indent, 0)
                # reset deeper levels when going up
                for k in list(list_counters):
                    if k > indent:
                        del list_counters[k]
                list_counters[indent] += 1
                bullet = f"{list_counters[indent]}. "
            else:
                bullet = "• "

            para = doc.add_paragraph()
            bullet_run = para.add_run(bullet)
            apply_run_format(bullet_run, body["font"], body["font_en"],
                             body["size_pt"], bold=ordered)
            add_inline(para, block["text"],
                       body["font"], body["font_en"], body["size_pt"])
            left = 0.5 + indent * 0.5
            set_paragraph_format(para, space_before=2, space_after=2,
                                 left_indent=left,
                                 first_line_indent=-Pt(body["size_pt"]).cm * 28)

        elif btype == "table":
            header = block["header"]
            rows = block["rows"]
            col_count = len(header)
            all_rows = [header] + [
                r + [""] * (col_count - len(r)) for r in rows
            ]
            border_mode = tbl_cfg.get("border_mode", "grid")
            rule_color = tbl_cfg.get("rule_color", "000000")
            tbl = doc.add_table(rows=len(all_rows), cols=col_count)
            tbl.style = "Table Grid"

            _THICK = {"val": "single", "sz": 12, "color": rule_color}
            _THIN  = {"val": "single", "sz": 6,  "color": rule_color}
            _NIL   = {"val": "nil"}
            accent_color = tbl_cfg.get("accent_color", rule_color)
            row_sep_color = tbl_cfg.get("row_sep_color", "E0E0E0")
            _ACCENT = {"val": "single", "sz": 8, "color": accent_color}
            _ROW_SEP = {"val": "single", "sz": 4, "color": row_sep_color}

            for row_idx, row_data in enumerate(all_rows):
                is_header = row_idx == 0
                is_last   = row_idx == len(all_rows) - 1
                tr = tbl.rows[row_idx]
                for col_idx, cell_text in enumerate(row_data):
                    cell = tr.cells[col_idx]

                    if border_mode == "grid":
                        c = tbl_cfg.get("border_color", "AAAAAA")
                        b = {"val": "single", "sz": 4, "color": c}
                        _set_cell_borders(cell, top=b, bottom=b, left=b, right=b)
                    elif border_mode == "mckinsey":
                        _set_cell_borders(
                            cell,
                            top=_THICK if is_header else _NIL,
                            bottom=_THIN if (is_header or is_last) else _NIL,
                            left=_NIL, right=_NIL,
                        )
                    elif border_mode == "rb":
                        # Header: thick black top + accent-color bottom
                        # Data rows: light gray separator between each row
                        # Last row: thin black bottom
                        _set_cell_borders(
                            cell,
                            top=_THICK if is_header else _NIL,
                            bottom=_ACCENT if is_header else (_THIN if is_last else _ROW_SEP),
                            left=_NIL, right=_NIL,
                        )

                    if is_header and tbl_cfg.get("header_bg_color"):
                        _set_cell_shading(cell, tbl_cfg["header_bg_color"])
                    para = cell.paragraphs[0]
                    para.clear()
                    add_inline(para, cell_text,
                               tbl_cfg["font"], tbl_cfg["font_en"],
                               tbl_cfg["size_pt"],
                               base_bold=is_header and tbl_cfg.get("header_bold", True))
                    set_paragraph_format(para,
                                         space_before=tbl_cfg.get("cell_padding_pt", 4),
                                         space_after=tbl_cfg.get("cell_padding_pt", 4))
            doc.add_paragraph()

    return doc


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to DOCX with custom formatting.")
    parser.add_argument("input", nargs="?", help="Input .md file")
    parser.add_argument("output", nargs="?", help="Output .docx file (default: same name)")
    parser.add_argument("--style", help="JSON style config file")
    parser.add_argument("--dump-style", action="store_true",
                        help="Print default style JSON and exit")
    args = parser.parse_args()

    if args.dump_style:
        _style_file = ASSETS_DIR / "default-style.json"
        if not _style_file.exists():
            print(f"Error: {_style_file} not found — reinstall the skill", file=sys.stderr)
            sys.exit(1)
        print(_style_file.read_text(encoding="utf-8"))
        return

    if not args.input:
        parser.error("input is required unless --dump-style is used")
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")

    _style_file = ASSETS_DIR / "default-style.json"
    if not _style_file.exists():
        print(f"Error: {_style_file} not found — reinstall the skill", file=sys.stderr)
        sys.exit(1)
    with open(_style_file, encoding="utf-8") as f:
        style = json.load(f)
    if args.style:
        with open(args.style, encoding="utf-8") as f:
            user_style = json.load(f)
        for section, values in user_style.items():
            if section in style and isinstance(style[section], dict):
                if section == "headings":
                    for h, hvals in values.items():
                        style[section].setdefault(h, {}).update(hvals)
                else:
                    style[section].update(values)
            else:
                style[section] = values

    md_text = input_path.read_text(encoding="utf-8")
    blocks = parse_md_blocks(md_text)
    doc = build_docx(blocks, style, base_dir=input_path.resolve().parent)
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
